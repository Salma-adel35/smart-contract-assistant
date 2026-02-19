import fitz
import docx
from langchain_core.documents import Document

def parse_files(uploaded_files):
    documents = []

    for file in uploaded_files:

        if file.name.endswith(".pdf"):
            with fitz.open(stream=file.read(), filetype="pdf") as pdf:
                text = ""
                for page in pdf:
                    text += page.get_text()

        elif file.name.endswith(".docx"):
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])

        else:
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={"source": file.name}
            )
        )

    return documents
